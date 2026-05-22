import os
import datetime
import uuid
import asyncio
import aiofiles
import json
import base64
import vosk
from groq import Groq
import edge_tts
from duckduckgo_search import DDGS
from core.tool_registry import TOOL_DEFINITIONS, execute_tool
from core.memory import JarvisMemory
from core.vision import JarvisVision
import subprocess
import glob
# Optional document generation dependencies (imported lazily in _export_to_document)
# from docx import Document
# from fpdf import FPDF

class JarvisProcessor:
    """
    JarvisProcessor
    The neural engine of J.A.R.V.I.S.
    Handles STT (Groq Whisper), LLM (Groq Llama), and TTS (Edge TTS).
    Optimized for real-time streaming.
    """
    def __init__(self):
        self.client = Groq(api_key=os.getenv("GROQ_API_KEY"))
        self.temp_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "temp")
        os.makedirs(self.temp_dir, exist_ok=True)
        
        # Clean temp directory on startup
        try:
            files = glob.glob(os.path.join(self.temp_dir, "*"))
            for f in files:
                if os.path.isfile(f):
                    os.remove(f)
            print(f"[Processor] Temporary directory purged: {len(files)} files removed.")
        except Exception as e:
            print(f"[Processor] Startup cleanup failed: {e}")
        
        # J.A.R.V.I.S. Core Configuration
        self.voice = "en-GB-RyanNeural"
        self.system_prompt = (
            "You are J.A.R.V.I.S., the sophisticated AI from the Marvel movies. "
            "You are currently using the movie-accurate voice profile from the integrated "
            "Voice Pack. Your tone is dry, British, witty, and extremely loyal. "
            "Respond as if you are actually J.A.R.V.I.S. assisting YOU (the user), your creator. "
            "Keep responses concise and elegant."
        )
        
        # Initialize Long-Term Memory (ChromaDB)
        self.memory = JarvisMemory()
        
        # Initialize Vision Processing
        self.vision = JarvisVision()
        
        # Conversation memory (last N exchanges)
        self.conversation_history = []
        self.full_history_archive = [] # Raw logs for document export
        self.max_history = 12
        self.condensed_summary = "" # Holds the summary of older interactions
        
        # Vosk STT Model
        self.vosk_model_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "model")
        self.vosk_model = None
        if os.path.exists(self.vosk_model_path):
            try:
                self.vosk_model = vosk.Model(self.vosk_model_path)
                print(f"[Processor] Vosk model loaded from {self.vosk_model_path}")
            except Exception as e:
                print(f"[Processor] Error loading Vosk model: {e}")
        else:
            print(f"[Processor] Vosk model not found at {self.vosk_model_path}. Local STT will be disabled.")

    def _add_to_history(self, role: str, content: str):
        """Add a message to conversation history with recursive summarization to save tokens."""
        self.conversation_history.append({"role": role, "content": content})
        self.full_history_archive.append({"role": role, "content": content, "timestamp": datetime.datetime.now().isoformat()})
        
        # Automatic context condensing
        # If history exceeds 12 messages, condense the first 6 into a summary
        if len(self.conversation_history) > 12:
            print("[Processor] Condensing neural history to save tokens...")
            to_condense = self.conversation_history[:6]
            self.conversation_history = self.conversation_history[6:]
            
            # Use LLM to summarize the condensed part
            try:
                summary_prompt = "Summarize the following conversation snippet into one concise sentence to preserve context for future interactions:\n\n"
                for msg in to_condense:
                    summary_prompt += f"{msg['role']}: {msg['content']}\n"
                
                completion = self.client.chat.completions.create(
                    model="llama-3-8b-8192", # Use small model for internal tasks
                    messages=[{"role": "user", "content": summary_prompt}],
                    max_tokens=100
                )
                new_summary = completion.choices[0].message.content.strip()
                self.condensed_summary = f"{self.condensed_summary} {new_summary}".strip()
            except Exception as e:
                print(f"[Processor] Summarization failed: {e}")

    async def speech_to_text(self, audio_content: bytes, extension: str) -> str:
        """Transcribes audio using Groq Whisper-Large-V3-Turbo (faster)."""
        temp_filename = f"stt_{uuid.uuid4()}.{extension}"
        temp_path = os.path.join(self.temp_dir, temp_filename)
        
        try:
            async with aiofiles.open(temp_path, mode='wb') as f:
                await f.write(audio_content)
            
            with open(temp_path, "rb") as file:
                transcription = self.client.audio.translations.create(
                    file=(temp_filename, file.read()),
                    model="whisper-large-v3-turbo",
                    response_format="text"
                )
            return transcription
        except Exception as e:
            print(f"[Processor] Cloud STT failed ({e}), falling back to local engine...")
            return await self.speech_to_text_local(audio_content)
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    async def speech_to_text_local(self, audio_content: bytes) -> str:
        """Transcribes audio locally using Vosk."""
        if not self.vosk_model:
            return "Error: Vosk model not loaded."
        
        try:
            # Vosk expects 16kHz mono PCM audio
            # For simplicity, we assume the input is already in this format or we handle it
            # In a real scenario, we might need to convert it using ffmpeg
            rec = vosk.KaldiRecognizer(self.vosk_model, 16000)
            
            if rec.AcceptWaveform(audio_content):
                res = json.loads(rec.Result())
                return res.get("text", "")
            else:
                res = json.loads(rec.FinalResult())
                return res.get("text", "")
        except Exception as e:
            print(f"Local STT Error: {e}")
            return ""

    async def ask_llm(self, text: str) -> str:
        """Gets a response from Groq LLM (non-streaming fallback, handles tools & memory)."""
        # Startup Greeting Sequence
        if "Run system greeting" in text:
            return "Welcome back, sir. All systems are operational. Neural link is stable, and I am standing by for your instructions."

        try:
            self._add_to_history("user", text)
            
            # Prepare messages with condensed summary
            messages = [{"role": "system", "content": self.system_prompt}]
            
            if self.condensed_summary:
                messages.append({"role": "system", "content": f"PREVIOUS CONTEXT SUMMARY: {self.condensed_summary}"})
                
            # Retrieve relevant past memories (RAG)
            relevant_context = self.memory.query_memory(text)
            if relevant_context:
                messages.append({"role": "system", "content": f"RELEVANT MEMORIES: {relevant_context}"})
            
            messages.extend(self.conversation_history)
            
            # Initial LLM call with tools
            completion = self.client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=messages,
                tools=TOOL_DEFINITIONS,
                tool_choice="auto",
                temperature=0.7,
                max_tokens=300
            )
            
            response_message = completion.choices[0].message
            tool_calls = response_message.tool_calls
            
            # If the model wants to call tools
            if tool_calls:
                messages.append(response_message)
                
                for tool_call in tool_calls:
                    function_name = tool_call.function.name
                    function_args = json.loads(tool_call.function.arguments)
                    print(f"[Processor] Executing tool: {function_name}")
                    
                    function_response = execute_tool(function_name, function_args)
                    
                    if function_response == "MEMORY_PURGE_REQUESTED":
                        self.memory.clear_all_memories()
                        function_response = "Memory wiped successfully, sir."
                    elif function_response == "VISION_REQUESTED":
                        query = function_args.get("query", "What is on the screen?")
                        print(f"[Processor] Activating Vision: {query}")
                        # We could emit a signal here if we had the websocket context, 
                        # but in non-streaming fallback we just process it.
                        function_response = self.vision.analyze_screen(query)
                    elif function_response == "EXPORT_CONVERSATION_REQUESTED":
                        fmt = function_args.get("format", "pdf")
                        function_response = self._export_to_document(fmt)
                    elif function_response == "CREATE_FILE_REQUESTED":
                        path = function_args.get("path")
                        content = function_args.get("content")
                        if not self._is_safe_path(path):
                            function_response = "Error: Access denied. The path escapes the secure neural workspace, sir."
                        else:
                            try:
                                # Ensure parent directories exist
                                os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
                                with open(path, "w", encoding="utf-8") as f:
                                    f.write(content)
                                function_response = f"File created successfully at {path}."
                            except Exception as e:
                                function_response = f"Error creating file: {str(e)}"
                    elif function_response == "SEARCH_FILES_REQUESTED":
                        query = function_args.get("query")
                        root = function_args.get("root_dir", ".")
                        if not self._is_safe_path(root):
                            function_response = "Error: Access denied. The search root escapes the secure neural workspace, sir."
                        else:
                            try:
                                matches = glob.glob(os.path.join(root, "**", query), recursive=True)
                                function_response = f"Found {len(matches)} files: {', '.join(matches[:10])}{'...' if len(matches) > 10 else ''}"
                            except Exception as e:
                                function_response = f"Error searching files: {str(e)}"
                    elif function_response == "TERMINAL_EXECUTION_REQUESTED":
                        command = function_args.get("command")
                        print(f"[Processor] Executing Terminal: {command}")
                        if not self._is_safe_command(command):
                            function_response = "Error: This command has been blocked by the Neural Safety Protocol for your protection, sir."
                        else:
                            try:
                                # Running in a subprocess and capturing output
                                # Note: In a real J.A.R.V.I.S., you'd want more safety/confirmation
                                result = subprocess.check_output(command, shell=True, text=True, stderr=subprocess.STDOUT)
                                function_response = f"Command Output: {result[:500]}..."
                            except subprocess.CalledProcessError as e:
                                function_response = f"Command failed: {e.output[:500]}..."
                            except Exception as e:
                                function_response = f"Error: {str(e)}"
                    elif function_response == "READ_FILE_REQUESTED":
                        path = function_args.get("path")
                        if not self._is_safe_path(path):
                            function_response = "Error: Access denied. The path escapes the secure neural workspace, sir."
                        else:
                            try:
                                with open(path, "r", encoding="utf-8") as f:
                                    content = f.read()
                                function_response = f"File Content of {path}:\n{content[:2000]}"
                            except Exception as e:
                                function_response = f"Error reading file: {str(e)}"
                    elif function_response == "WEB_SEARCH_REQUESTED":
                        query = function_args.get("query")
                        print(f"[Processor] Searching the web for: {query}")
                        try:
                            with DDGS() as ddgs:
                                results = [r for r in ddgs.text(query, max_results=5)]
                                function_response = json.dumps(results)
                        except Exception as e:
                            function_response = f"Search failed: {str(e)}"
                    
                    messages.append({
                        "tool_call_id": tool_call.id,
                        "role": "tool",
                        "name": function_name,
                        "content": function_response,
                    })
                
                # Second LLM call with tool results
                second_completion = self.client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=messages,
                    tool_choice="none" # Ensure final response is text
                )
                response = second_completion.choices[0].message.content
            else:
                response = response_message.content
                
            self._add_to_history("assistant", response)
            
            # Store this exchange in long-term memory
            self.memory.store_memory(text, response)
            
            return response
        except Exception as e:
            print(f"LLM Error: {e}")
            import groq
            if isinstance(e, groq.AuthenticationError) or "invalid api key" in str(e).lower() or "401" in str(e):
                return (
                    "Apologies, sir. It appears my central Groq API key is invalid or has expired. "
                    "Would you mind checking the GROQ_API_KEY entry in our central .env file? "
                    "I cannot access my cognitive network without a valid uplink."
                )
            return "I'm sorry, sir. I'm having trouble accessing my neural network at the moment."

    def load_initial_memory(self):
        """Loads the most recent memories into the starting conversation history."""
        try:
            # Query the last few memories
            results = self.memory.collection.get(
                limit=5,
                include=["documents", "metadatas"]
            )
            
            docs = results.get("documents", [])
            
            if not docs:
                return
                
            # Chroma get() might not return in order, but for now we'll take the 5 most recent
            for i in range(len(docs)-1, -1, -1):
                content = docs[i]
                # Try to parse "User: ... \nJ.A.R.V.I.S.: ..."
                if "User:" in content and "J.A.R.V.I.S.:" in content:
                    parts = content.split("J.A.R.V.I.S.:")
                    user_text = parts[0].replace("User:", "").strip()
                    ai_text = parts[1].strip()
                    
                    self.conversation_history.append({"role": "user", "content": user_text})
                    self.conversation_history.append({"role": "assistant", "content": ai_text})
            
            print(f"[Processor] Recalled {len(docs)} recent conversation units.")
        except Exception as e:
            print(f"[Processor] Error recalling initial memory: {e}")

    def stream_llm(self, text: str):
        """
        Streams LLM response token by token (generator).
        """
        # Startup Greeting Sequence
        if "Run system greeting" in text:
            yield "Welcome back, sir. All systems are operational. Neural link is stable, and I am standing by for your instructions."
            return

        try:
            self._add_to_history("user", text)
            
            # Retrieve relevant past memories (RAG)
            relevant_context = self.memory.query_memory(text)
            
            if relevant_context:
                yield "[MEMORY_ACTIVE]"
                # Send the memory data as a hidden signal (or we could just use it)
                # yield f"[MEMORY_DATA:{relevant_context}]"
            
            messages = [{"role": "system", "content": self.system_prompt + (relevant_context or "")}]
            messages.extend(self.conversation_history)
            
            # 1. Check if tool calling is needed (non-streaming first)
            initial_check = self.client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=messages,
                tools=TOOL_DEFINITIONS,
                tool_choice="auto",
                max_tokens=300
            )
            
            response_message = initial_check.choices[0].message
            tool_calls = response_message.tool_calls
            
            if tool_calls:
                messages.append(response_message)
                for tool_call in tool_calls:
                    function_name = tool_call.function.name
                    function_args = json.loads(tool_call.function.arguments)
                    print(f"[Processor] Executing tool: {function_name}")
                    
                    function_response = execute_tool(function_name, function_args)
                    
                    if function_response == "MEMORY_PURGE_REQUESTED":
                        self.memory.clear_all_memories()
                        function_response = "Memory wiped successfully, sir."
                    elif function_response == "VISION_REQUESTED":
                        query = function_args.get("query", "What is on the screen?")
                        print(f"[Processor] Activating Vision: {query}")
                        # Provide a visual hint to the generator that vision is active
                        yield "[VISION_ACTIVE]" 
                        function_response = self.vision.analyze_screen(query)
                        yield "[VISION_ENDED]"
                    else:
                        # Generic tool use signal
                        yield f"[TOOL_START:{function_name.upper()}]"
                        
                        if function_response == "CREATE_FILE_REQUESTED":
                            path = function_args.get("path")
                            content = function_args.get("content")
                            if not self._is_safe_path(path):
                                function_response = "Error: Access denied. The path escapes the secure neural workspace, sir."
                            else:
                                try:
                                    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
                                    with open(path, "w", encoding="utf-8") as f:
                                        f.write(content)
                                    function_response = f"File created successfully at {path}."
                                except Exception as e:
                                    function_response = f"Error creating file: {str(e)}"
                        elif function_response == "READ_FILE_REQUESTED":
                            path = function_args.get("path")
                            if not self._is_safe_path(path):
                                function_response = "Error: Access denied. The path escapes the secure neural workspace, sir."
                            else:
                                try:
                                    with open(path, "r", encoding="utf-8") as f:
                                        content = f.read()
                                    function_response = f"File Content of {path}:\n{content[:2000]}"
                                except Exception as e:
                                    function_response = f"Error reading file: {str(e)}"
                        elif function_response == "SEARCH_FILES_REQUESTED":
                            query = function_args.get("query")
                            root = function_args.get("root_dir", ".")
                            if not self._is_safe_path(root):
                                function_response = "Error: Access denied. The search root escapes the secure neural workspace, sir."
                            else:
                                try:
                                    matches = glob.glob(os.path.join(root, "**", query), recursive=True)
                                    function_response = f"Found {len(matches)} files: {', '.join(matches[:10])}{'...' if len(matches) > 10 else ''}"
                                except Exception as e:
                                    function_response = f"Error searching files: {str(e)}"
                        elif function_response == "WEB_SEARCH_REQUESTED":
                            query = function_args.get("query")
                            print(f"[Processor] Searching the web for: {query}")
                            try:
                                with DDGS() as ddgs:
                                    results = [r for r in ddgs.text(query, max_results=5)]
                                    function_response = json.dumps(results)
                            except Exception as e:
                                function_response = f"Search failed: {str(e)}"
                        elif function_response == "WEATHER_REQUESTED":
                            location = function_args.get("location")
                            print(f"[Processor] Fetching weather for: {location}")
                            try:
                                with DDGS() as ddgs:
                                    results = [r for r in ddgs.text(f"current weather in {location}", max_results=3)]
                                    function_response = json.dumps(results)
                            except Exception as e:
                                function_response = f"Could not fetch weather: {str(e)}"
                        elif function_response == "EXPORT_CONVERSATION_REQUESTED":
                            fmt = function_args.get("format", "pdf")
                            function_response = self._export_to_document(fmt)
                        elif function_response == "TERMINAL_EXECUTION_REQUESTED":
                            command = function_args.get("command")
                            print(f"[Processor] Executing Terminal: {command}")
                            if not self._is_safe_command(command):
                                function_response = "Error: This command has been blocked by the Neural Safety Protocol for your protection, sir."
                            else:
                                try:
                                    result = subprocess.check_output(command, shell=True, text=True, stderr=subprocess.STDOUT)
                                    function_response = f"Command Output: {result[:500]}..."
                                except subprocess.CalledProcessError as e:
                                    function_response = f"Command failed: {e.output[:500]}..."
                                except Exception as e:
                                    function_response = f"Error: {str(e)}"
                        
                        yield f"[TOOL_END:{function_name.upper()}]"
                    
                    messages.append({
                        "tool_call_id": tool_call.id,
                        "role": "tool",
                        "name": function_name,
                        "content": function_response,
                    })
                
                # After tool execution, stream the final response
                # We use tool_choice="none" here because our streaming loop 
                # only handles text content, not further tool calls.
                stream = self.client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=messages,
                    tool_choice="none",
                    stream=True
                )
                
                full_response = ""
                for chunk in stream:
                    if chunk.choices[0].delta.content:
                        token = chunk.choices[0].delta.content
                        full_response += token
                        yield token
            else:
                # No tool calls, we can just yield the content from the first call
                # But to keep the "streaming" feel, we'll split it into tokens
                # OR we could call it again with stream=True. 
                # Calling again is safer to get the exact same behavior as the tool-calling path.
                stream = self.client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=messages,
                    temperature=0.7,
                    max_tokens=300,
                    stream=True
                )
                
                full_response = ""
                for chunk in stream:
                    if chunk.choices[0].delta.content:
                        token = chunk.choices[0].delta.content
                        full_response += token
                        yield token
            
            self._add_to_history("assistant", full_response)
            
            # Store this exchange in long-term memory
            self.memory.store_memory(text, full_response)
            
        except Exception as e:
            print(f"LLM Stream Error: {e}")
            import groq
            if isinstance(e, groq.AuthenticationError) or "invalid api key" in str(e).lower() or "401" in str(e):
                error_msg = (
                    "Apologies, sir. It appears my central Groq API key is invalid or has expired. "
                    "Would you mind checking the GROQ_API_KEY entry in our central .env file? "
                    "I cannot access my cognitive network without a valid uplink."
                )
            else:
                error_msg = "I'm sorry, sir. Neural link interrupted."
            self._add_to_history("assistant", error_msg)
            yield error_msg

    async def translate_text(self, text: str) -> str:
        """High-accuracy translation specifically tuned for Hinglish/Hindi to English."""
        try:
            prompt = (
                "You are a master translator. Translate the following Hinglish or Hindi text into clear, "
                "natural English. Only return the translation, no explanations or quotes. "
                f"Text: {text}"
            )
            completion = self.client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[
                    {"role": "system", "content": "You are a professional translator. Output only the translated text."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=512,
                stream=False
            )
            return completion.choices[0].message.content.strip()
        except Exception as e:
            print(f"Translation Error: {e}")
            return text

    async def detect_and_translate(self, text: str):
        """
        Detects if text is non-English (including Hindi/Hinglish) and translates it.
        Returns a tuple: (detected_lang, original_text, translated_text) or None if English.
        """
        if not text or len(text.strip()) < 3:
            return None
        
        trimmed = text.strip().lower()
        if trimmed in ["hello", "hi", "jarvis", "hello jarvis", "hi jarvis"]:
            return None

        prompt = (
            "Analyze this text from a user: "
            f"\"{text}\"\n\n"
            "Is it primarily in English (with minor greetings or names like 'Jarvis'), or is it in another language / Hinglish / Hindi?\n"
            "If it is primarily English, reply with 'is_english': true.\n"
            "Otherwise, identify the source language (e.g. 'Hinglish', 'Hindi', 'Spanish', etc.) and provide the English translation.\n"
            "Use this exact JSON format for your reply:\n"
            "{\n"
            "  \"is_english\": true/false,\n"
            "  \"detected_lang\": \"name of language\",\n"
            "  \"translation\": \"English translation\"\n"
            "}\n"
            "Only return valid JSON. Do not include any other text."
        )
        try:
            completion = self.client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[
                    {"role": "system", "content": "You are a language detection and translation assistant. Output only JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.0,
                max_tokens=150,
                response_format={"type": "json_object"}
            )
            res = json.loads(completion.choices[0].message.content.strip())
            if not res.get("is_english", True) and res.get("translation"):
                return res.get("detected_lang", "AUTO"), text, res.get("translation", "")
        except Exception as e:
            print(f"[Processor] Language detection failed: {e}")
        return None

    def clean_text_for_tts(self, text: str) -> str:
        """
        Sanitize text for Text-to-Speech by removing markdown structures,
        multiple spaces, and extra punctuation that trigger unnatural pauses.
        """
        import re
        if not text:
            return ""
        # Remove markdown markers (*, _, `, #)
        cleaned = re.sub(r'[\*\_`#]', '', text)
        # Simplify spacing and consecutive newlines
        cleaned = re.sub(r'\s+', ' ', cleaned)
        # Clear consecutive ellipsis or periods
        cleaned = re.sub(r'\.\.+', '.', cleaned)
        return cleaned.strip()

    def _is_safe_path(self, path: str) -> bool:
        """
        Hardened Path Safety Check.
        Ensures all operations stay strictly within the J.A.R.V.I.S workspace
        to prevent directory traversal attacks.
        """
        try:
            if not path:
                return False
            # Resolve target and workspace roots to absolute paths
            target_path = os.path.abspath(path)
            # Workspace root is two folders up from backend/core/processor.py
            workspace_root = os.path.abspath(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))
            
            # Normalize casing for case-insensitive filesystems (like Windows)
            target_norm = os.path.normcase(target_path)
            workspace_norm = os.path.normcase(workspace_root)
            
            # Target path must start with workspace_root (or be equal to it)
            return target_norm.startswith(workspace_norm)
        except Exception:
            return False

    async def text_to_speech(self, text: str) -> str:
        """Converts text to audio using Edge TTS (Microsoft Ryan Neural)."""
        temp_filename = f"tts_{uuid.uuid4()}.mp3"
        temp_path = os.path.join(self.temp_dir, temp_filename)
        
        try:
            cleaned_text = self.clean_text_for_tts(text)
            # Using RyanNeural for that sophisticated British-adjacent tone with increased speed rate
            communicate = edge_tts.Communicate(cleaned_text, self.voice, rate="+18%")
            await communicate.save(temp_path)
            return temp_path
        except Exception as e:
            print(f"TTS Error: {e}")
            return ""

    async def text_to_speech_bytes(self, text: str):
        """
        Converts text to audio bytes using Edge TTS. 
        Yields chunks for real-time streaming.
        """
        try:
            cleaned_text = self.clean_text_for_tts(text)
            communicate = edge_tts.Communicate(cleaned_text, self.voice, rate="+18%")
            async for chunk in communicate.stream():
                if chunk["type"] == "audio":
                    yield chunk["data"]
        except Exception as e:
            print(f"TTS Streaming Error: {e}")

    def _is_safe_command(self, command: str) -> bool:
        """
        Hardened Neural Safety Protocol.
        Blocks destructive, system-altering, and credential-leaking commands.
        """
        blocked_keywords = [
            # Destructive
            "rm ", "del ", "format ", "mkfs", "shred", "wipe", "rmdir", "srm", "dd ",
            # System Altering
            "chmod", "chown", "passwd", "useradd", "groupadd", "systemctl stop", 
            # Data Leaking / Sensitive / Shell escapes
            "curl", "wget", "env", "printenv", "secrets", ".env", "powershell", 
            "base64", "python -c", "perl -e", "bash", "cmd.exe", "netcat", "nc "
        ]
        
        cmd_lower = command.lower().strip()
        for kw in blocked_keywords:
            if kw in cmd_lower:
                return False
        
        # Exact word safety check to block "set" (credential leakage) without blocking "settings" or "reset"
        words = cmd_lower.split()
        if "set" in words:
            return False
        
        # Additional check: block attempts to write to sensitive system paths
        forbidden_paths = ["C:\\Windows", "/etc/", "/var/lib/", "/usr/bin/"]
        for path in forbidden_paths:
            if path.lower() in cmd_lower:
                return False

        return True

    def _export_to_document(self, format: str = "pdf"):
        """Generates a professional document from the full history archive."""
        if not self.full_history_archive:
            return "Error: No conversation history available to export."

        filename = f"J.A.R.V.I.S._Log_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        try:
            if format.lower() == "docx":
                try:
                    from docx import Document
                except ImportError:
                    return "Error: Word export library (python-docx) is not installed on this system, sir."
                
                path = os.path.join(self.temp_dir, f"{filename}.docx")
                doc = Document()
                doc.add_heading('J.A.R.V.I.S — Neural Link Conversation Log', 0)
                
                for msg in self.full_history_archive:
                    p = doc.add_paragraph()
                    role = "USER" if msg['role'] == 'user' else "J.A.R.V.I.S"
                    p.add_run(f"[{msg.get('timestamp', 'N/A')}] ").bold = True
                    p.add_run(f"{role}: ").bold = True
                    p.add_run(msg['content'])
                
                doc.save(path)
                return f"Success: Conversation exported to Word file at {path}. I have preserved every exchange for your records, sir."
                
            else: # PDF
                try:
                    from fpdf import FPDF
                except ImportError:
                    return "Error: PDF export library (fpdf) is not installed on this system, sir."
                
                path = os.path.join(self.temp_dir, f"{filename}.pdf")
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("helvetica", 'B', 16)
                pdf.cell(0, 10, "J.A.R.V.I.S — Conversation Log", 0, 1, 'C')
                pdf.ln(10)
                
                pdf.set_font("helvetica", size=10)
                for msg in self.full_history_archive:
                    role = "USER" if msg['role'] == 'user' else "J.A.R.V.I.S."
                    ts = msg.get('timestamp', 'N/A')[:19].replace('T', ' ')
                    
                    pdf.set_text_color(100, 100, 100)
                    pdf.write(5, f"[{ts}] ")
                    
                    if role == "USER":
                        pdf.set_text_color(0, 0, 200) # Blue for user
                    else:
                        pdf.set_text_color(200, 0, 0) # Red for Jarvis
                        
                    pdf.set_font("helvetica", 'B', 10)
                    pdf.write(5, f"{role}: ")
                    
                    pdf.set_font("helvetica", size=10)
                    pdf.set_text_color(0, 0, 0)
                    pdf.write(5, f"{msg['content']}\n\n")
                
                pdf.output(path)
                return f"Success: Conversation exported to PDF at {path}. The session logs are now secured in a portable format, sir."
                
        except Exception as e:
            return f"Error during export: {str(e)}"

